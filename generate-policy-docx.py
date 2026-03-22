#!/usr/bin/env python3
"""
generate-policy-docx.py — Converts Open IT Policy Library markdown files to Word (.docx)

Usage:
    python generate-policy-docx.py AI-Acceptable-Use-Policy-AEC.md
    python generate-policy-docx.py *.md

Reads a policy markdown file and generates a professionally formatted .docx
with headers, footers, branded colors, tables, and signature blocks.

Works with any policy following the Open IT Policy Library formatting conventions.
"""

import sys
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

# ── Brand colors ──────────────────────────────────────────────────────────────
BRAND_BLUE = RGBColor(0x2B, 0x57, 0x9A)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
TABLE_HEADER_BG = "2B579A"
TABLE_ALT_ROW = "F2F2F2"
BORDER_COLOR = "BFBFBF"
ADMONITION_TIP_BG = "E8F4E8"       # Light green background for tip admonitions
ADMONITION_TIP_BORDER = "2E7D32"   # Green left border for tips
ADMONITION_WARN_BG = "FFF3E0"      # Light orange background for warning admonitions
ADMONITION_WARN_BORDER = "E65100"  # Orange left border for warnings

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_bottom_border(paragraph, color="D6E4F0", size=6):
    """Add a bottom border to a paragraph (used as a horizontal rule).

    Schema order for pPr children:
    pStyle, keepNext, keepLines, pageBreakBefore, ..., numPr, ..., pBdr, shd, tabs,
    ..., spacing, ind, ..., jc, ..., rPr

    pBdr must come after pStyle/numPr but before shd/tabs/spacing/ind/jc/rPr.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    # Find correct insertion point: after pStyle, numPr, etc. but before shd, tabs, spacing, ind, jc, rPr
    later_elements = ['shd', 'tabs', 'suppressAutoHyphens', 'spacing', 'ind', 'jc', 'rPr']
    insert_before = None
    for child in pPr:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in later_elements:
            insert_before = child
            break
    if insert_before is not None:
        insert_before.addprevious(pBdr)
    else:
        pPr.append(pBdr)


def set_run_font(run, font_name="Arial", size=None, bold=False, italic=False, color=None):
    """Configure a run's font properties."""
    run.font.name = font_name
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_hyperlink(paragraph, text, url, font_size=8, color=None, bold=False, italic=False):
    """Add a hyperlink run to a paragraph."""
    link_color = color or BRAND_BLUE
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}></w:hyperlink>')
    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr>'
        f'    <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
        f'    <w:sz w:val="{int(font_size * 2)}"/>'
        f'    <w:color w:val="{link_color}"/>'
        f'    <w:u w:val="single"/>'
        f'    {"<w:b/>" if bold else ""}'
        f'    {"<w:i/>" if italic else ""}'
        f'  </w:rPr>'
        f'  <w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_formatted_text_with_links(paragraph, text, font_size=10.5, color=DARK_GRAY, bold=False, italic=False):
    """Add formatted text to a paragraph, rendering markdown links as clickable hyperlinks."""
    # Split on markdown links: [text](url)
    parts = re.split(r'(\[[^\]]+\]\([^)]+\))', text)
    for part in parts:
        link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
        if link_match:
            link_text = link_match.group(1)
            link_url = link_match.group(2)
            add_hyperlink(paragraph, link_text, link_url, font_size=font_size, color=BRAND_BLUE, bold=bold, italic=italic)
        else:
            # Handle **bold** markers within the text
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bp in bold_parts:
                if bp.startswith('**') and bp.endswith('**'):
                    run = paragraph.add_run(bp[2:-2])
                    set_run_font(run, size=font_size, bold=True, italic=italic, color=color)
                else:
                    run = paragraph.add_run(bp)
                    set_run_font(run, size=font_size, bold=bold, italic=italic, color=color)


def add_formatted_text(paragraph, text, font_size=10.5, color=DARK_GRAY, bold=False, italic=False):
    """Add a run with formatting to a paragraph, handling **bold** markers."""
    # Split on **bold** markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=font_size, bold=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=font_size, bold=bold, italic=italic, color=color)


def add_horizontal_rule(doc):
    """Add a styled horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    add_bottom_border(p)
    return p


# ── Markdown parser ───────────────────────────────────────────────────────────

def parse_markdown(filepath):
    """Parse a policy markdown file into structured blocks."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Strip <br> tags — they're for HTML rendering, not DOCX
        line = re.sub(r'<br\s*/?>', '', line)

        # Strip "Download DOCX version" links — not needed inside the DOCX itself
        if 'Download DOCX version' in line:
            line = re.sub(r'\s*\|\s*\[Download DOCX version\]\([^)]*\)', '', line)
            line = line.rstrip()
            if not line.strip():
                i += 1
                continue

        # Horizontal rule
        if line.strip() == '---':
            blocks.append({'type': 'hr'})
            i += 1
            continue

        # H1: # Title
        if line.startswith('# ') and not line.startswith('## '):
            blocks.append({'type': 'h1', 'text': line[2:].strip()})
            i += 1
            continue

        # H2: ## Section
        if line.startswith('## ') and not line.startswith('### '):
            blocks.append({'type': 'h2', 'text': line[3:].strip()})
            i += 1
            continue

        # H3: ### Subsection
        if line.startswith('### '):
            blocks.append({'type': 'h3', 'text': line[4:].strip()})
            i += 1
            continue

        # Table (| ... | ... |)
        if line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            # Parse table
            headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
            rows = []
            for tl in table_lines[2:]:  # skip header and separator
                row = [c.strip() for c in tl.split('|')[1:-1]]
                rows.append(row)
            blocks.append({'type': 'table', 'headers': headers, 'rows': rows})
            continue

        # Admonition (!!! type "Title")
        adm_match = re.match(r'^!!!\s+(\w+)\s+"([^"]+)"', line.strip())
        if adm_match:
            adm_type = adm_match.group(1)   # tip, warning, note, etc.
            adm_title = adm_match.group(2)
            adm_content_lines = []
            i += 1
            while i < len(lines):
                adm_line = lines[i].rstrip('\n')
                # Admonition content is indented by 4 spaces
                if adm_line.startswith('    '):
                    adm_content_lines.append(adm_line[4:])
                    i += 1
                elif adm_line.strip() == '':
                    # Empty line could be inside the admonition (between paragraphs/bullets)
                    # Check if next non-empty line is still indented
                    peek = i + 1
                    while peek < len(lines) and lines[peek].strip() == '':
                        peek += 1
                    if peek < len(lines) and lines[peek].startswith('    '):
                        adm_content_lines.append('')
                        i += 1
                    else:
                        break
                else:
                    break
            # Strip trailing empty lines
            while adm_content_lines and not adm_content_lines[-1].strip():
                adm_content_lines.pop()
            blocks.append({
                'type': 'admonition',
                'adm_type': adm_type,
                'title': adm_title,
                'text': '\n'.join(adm_content_lines)
            })
            continue

        # Blockquote (> ...)
        if line.strip().startswith('> '):
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith('> ') or lines[i].strip().startswith('>') and len(lines[i].strip()) <= 1):
                cleaned = lines[i].strip()
                if cleaned == '>':
                    quote_lines.append('')
                else:
                    quote_lines.append(cleaned[2:])
                i += 1
            blocks.append({'type': 'blockquote', 'text': '\n'.join(quote_lines)})
            continue

        # Numbered list (1. ...)
        num_match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if num_match:
            items = []
            while i < len(lines):
                stripped = lines[i].strip()
                m = re.match(r'^(\d+)\.\s+(.*)', stripped)
                if m:
                    items.append(m.group(2))
                    i += 1
                elif stripped == '':
                    break
                else:
                    # Continuation of previous item
                    if items:
                        items[-1] += ' ' + stripped
                    i += 1
            blocks.append({'type': 'numbered_list', 'items': items})
            continue

        # Bullet list (- ...)
        if line.strip().startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append({'type': 'bullet_list', 'items': items})
            continue

        # Bold + italic subtitle line (e.g., **Open IT Policy Library...** )
        if line.strip().startswith('**') and '|' in line:
            blocks.append({'type': 'subtitle', 'text': line.strip()})
            i += 1
            continue

        # Italic line starting with *AEC examples:* or *This policy* (check BEFORE version)
        if line.strip().startswith('*AEC') or line.strip().startswith('*This policy'):
            raw = line.strip()
            # Strip outer italic markers: *text* -> text
            # Handle pattern like *AEC examples:* rest of text
            raw = re.sub(r'^\*([^*]+)\*\s*', r'\1 ', raw)
            # Also strip any trailing * if the whole line was wrapped
            raw = raw.strip().rstrip('*').strip()
            blocks.append({'type': 'italic_para', 'text': raw})
            i += 1
            continue

        # Italic version line (e.g., *Version 1.0...*)
        if line.strip().startswith('*') and not line.strip().startswith('**') and line.strip().endswith('*'):
            blocks.append({'type': 'version', 'text': line.strip().strip('*')})
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            para_lines = [line.strip()]
            i += 1
            # Collect continuation lines (non-empty, non-special)
            while i < len(lines):
                next_line = lines[i].strip()
                if (next_line == '' or next_line.startswith('#') or next_line.startswith('|') or
                    next_line.startswith('- ') or next_line.startswith('> ') or next_line == '---' or
                    re.match(r'^\d+\.\s+', next_line)):
                    break
                para_lines.append(next_line)
                i += 1
            blocks.append({'type': 'paragraph', 'text': ' '.join(para_lines)})
            continue

        # Empty line
        i += 1

    return blocks


# ── Document builder ──────────────────────────────────────────────────────────

def build_docx(blocks, output_path):
    """Convert parsed blocks into a formatted .docx file."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = DARK_GRAY

    # ── Configure heading styles ──
    for level, size, color in [(1, 14, BRAND_BLUE), (2, 12, DARK_GRAY), (3, 11, DARK_GRAY)]:
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Arial'
        h_style.font.size = Pt(size)
        h_style.font.bold = True
        h_style.font.color.rgb = color
        h_style.paragraph_format.space_before = Pt(18 if level == 1 else 14)
        h_style.paragraph_format.space_after = Pt(10 if level == 1 else 8)

    # ── Extract title from first h1 for header ──
    title_text = ""
    for b in blocks:
        if b['type'] == 'h1':
            title_text = b['text']
            break

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    run1 = hp.add_run(title_text)
    set_run_font(run1, size=9, color=BRAND_BLUE)
    run2 = hp.add_run('\tOpen IT Policy Library')
    set_run_font(run2, size=8, italic=True, color=MED_GRAY)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_bottom_border(hp, color="2B579A", size=4)

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    add_hyperlink(fp, 'sortedsolution.com/oitpl', 'https://www.sortedsolution.com/oitpl', font_size=8, color=MED_GRAY)
    run2 = fp.add_run('\tPage ')
    set_run_font(run2, size=8, color=MED_GRAY)
    # Page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run3 = fp.add_run()
    run3._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run4 = fp.add_run()
    set_run_font(run4, size=8, color=MED_GRAY)
    run4._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5 = fp.add_run()
    run5._r.append(fldChar2)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── Process blocks ──
    is_first_h1 = True
    blockquote_bullet_mode = False  # Track if we're in a blockquote that has bullets

    for idx, block in enumerate(blocks):
        btype = block['type']

        if btype == 'h1':
            if is_first_h1:
                # Title block — big and bold
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(block['text'])
                set_run_font(run, size=20, bold=True, color=BRAND_BLUE)
                is_first_h1 = False
            else:
                p = doc.add_heading(block['text'], level=1)

        elif btype == 'subtitle':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            # Strip markdown bold markers and links for clean display
            text = block['text'].replace('**', '')
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            run = p.add_run(text)
            set_run_font(run, size=11, color=MED_GRAY)

        elif btype == 'version':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            add_formatted_text_with_links(p, block['text'], font_size=10, italic=True, color=MED_GRAY)

        elif btype == 'h2':
            doc.add_heading(block['text'], level=2)

        elif btype == 'h3':
            doc.add_heading(block['text'], level=3)

        elif btype == 'hr':
            add_horizontal_rule(doc)

        elif btype == 'paragraph':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Handle "quoted" intro text like "AI tools" includes:
            text = block['text']
            # Clean markdown links: [text](url) -> text
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            add_formatted_text(p, text)

        elif btype == 'italic_para':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            text = block['text']
            # Handle mixed *italic:* bold pattern for AEC examples
            if text.startswith('AEC examples:'):
                # Strip links for AEC examples
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
                run = p.add_run('AEC examples: ')
                set_run_font(run, size=10.5, italic=True, bold=True, color=DARK_GRAY)
                rest = text[len('AEC examples: '):]
                run2 = p.add_run(rest)
                set_run_font(run2, size=10.5, italic=True, color=DARK_GRAY)
            else:
                # Render with hyperlinks preserved
                add_formatted_text_with_links(p, text, font_size=10.5, italic=True, color=DARK_GRAY)

        elif btype == 'bullet_list':
            for item in block['items']:
                p = doc.add_paragraph(style='List Bullet')
                p.clear()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                # Clean markdown links
                item = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', item)
                add_formatted_text(p, item)

        elif btype == 'numbered_list':
            for item_idx, item in enumerate(block['items']):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                # Manual number prefix
                num_run = p.add_run(f"{item_idx + 1}. ")
                set_run_font(num_run, size=10.5, color=DARK_GRAY)
                item = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', item)
                add_formatted_text(p, item)

        elif btype == 'blockquote':
            text = block['text']
            # Check if blockquote contains bullet items
            bq_lines = text.split('\n')
            has_bullets = any(l.strip().startswith('- ') for l in bq_lines)

            if has_bullets:
                # Render non-bullet intro, then bullets
                for bq_line in bq_lines:
                    bq_line = bq_line.strip()
                    if not bq_line:
                        continue
                    if bq_line.startswith('- '):
                        p = doc.add_paragraph(style='List Bullet')
                        p.clear()
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.left_indent = Inches(0.75)
                        item_text = bq_line[2:]
                        item_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', item_text)
                        add_formatted_text(p, item_text, font_size=10, italic=True, color=MED_GRAY)
                    else:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.space_after = Pt(4)
                        bq_line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', bq_line)
                        add_formatted_text(p, bq_line, font_size=10, italic=True, color=MED_GRAY)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
                # Clean up double newlines
                text = re.sub(r'\n+', ' ', text).strip()
                add_formatted_text(p, text, font_size=10, italic=True, color=MED_GRAY)

        elif btype == 'admonition':
            adm_type = block.get('adm_type', 'tip')
            title = block.get('title', '')
            text = block['text']

            # Pick colors based on admonition type
            if adm_type == 'warning':
                bg_color = ADMONITION_WARN_BG
                border_color = ADMONITION_WARN_BORDER
                title_color = RGBColor(0xE6, 0x51, 0x00)
            else:  # tip, note, info, etc.
                bg_color = ADMONITION_TIP_BG
                border_color = ADMONITION_TIP_BORDER
                title_color = RGBColor(0x2E, 0x7D, 0x32)

            # Use a single-cell table to create a callout box
            adm_table = doc.add_table(rows=1, cols=1)
            adm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = adm_table.rows[0].cells[0]
            cell.text = ''
            set_cell_shading(cell, bg_color)

            # Apply left border to the cell
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{bg_color}"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{bg_color}"/>'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{bg_color}"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

            # Set table width
            tbl = adm_table._tbl
            tblPr = tbl.tblPr
            existing_w = tblPr.find(qn('w:tblW'))
            if existing_w is not None:
                tblPr.remove(existing_w)
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="9360" w:type="dxa"/>')
            tblPr.append(tblW)

            # Title paragraph
            title_p = cell.paragraphs[0]
            title_p.clear()
            title_p.paragraph_format.space_before = Pt(4)
            title_p.paragraph_format.space_after = Pt(6)
            title_run = title_p.add_run(title)
            set_run_font(title_run, size=10.5, bold=True, color=title_color)

            # Content — split into lines, handle bullets vs paragraphs
            adm_lines = text.split('\n')
            for adm_line in adm_lines:
                adm_line_stripped = adm_line.strip()
                if not adm_line_stripped:
                    continue
                if adm_line_stripped.startswith('- '):
                    bp = cell.add_paragraph()
                    bp.paragraph_format.space_before = Pt(2)
                    bp.paragraph_format.space_after = Pt(2)
                    bp.paragraph_format.left_indent = Inches(0.25)
                    bp.paragraph_format.first_line_indent = Inches(-0.15)
                    bullet_text = adm_line_stripped[2:]
                    bullet_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', bullet_text)
                    bullet_run = bp.add_run('• ')
                    set_run_font(bullet_run, size=10, color=DARK_GRAY)
                    add_formatted_text(bp, bullet_text, font_size=10, color=DARK_GRAY)
                else:
                    cp = cell.add_paragraph()
                    cp.paragraph_format.space_before = Pt(2)
                    cp.paragraph_format.space_after = Pt(2)
                    clean_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', adm_line_stripped)
                    add_formatted_text(cp, clean_text, font_size=10, color=DARK_GRAY)

            # Space after admonition
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        elif btype == 'table':
            headers = block['headers']
            rows = block['rows']
            num_cols = len(headers)

            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'

            # Set table width to full content area
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
                tbl.insert(0, tblPr)
            # Remove existing tblW if present, then add ours
            # Schema order: tblStyle, tblpPr, tblOverlap, bidiVisual, tblStyleRowBandSize,
            #   tblStyleColBandSize, tblW, jc, ...
            existing_w = tblPr.find(qn('w:tblW'))
            if existing_w is not None:
                tblPr.remove(existing_w)
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="9360" w:type="dxa"/>')
            # Insert after tblStyle if present, otherwise at proper position
            later_tbl = ['jc', 'tblCellSpacing', 'tblInd', 'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook']
            insert_before = None
            for child in tblPr:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in later_tbl:
                    insert_before = child
                    break
            if insert_before is not None:
                insert_before.addprevious(tblW)
            else:
                tblPr.append(tblW)

            # Header row
            hdr_row = table.rows[0]
            for ci, header_text in enumerate(headers):
                cell = hdr_row.cells[ci]
                cell.text = ''
                set_cell_shading(cell, TABLE_HEADER_BG)
                p = cell.paragraphs[0]
                run = p.add_run(header_text)
                set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

            # Data rows
            for ri, row_data in enumerate(rows):
                data_row = table.rows[ri + 1]
                for ci in range(num_cols):
                    cell = data_row.cells[ci]
                    cell_text = row_data[ci] if ci < len(row_data) else ''
                    cell.text = ''
                    # Clean markdown links
                    cell_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cell_text)
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    set_run_font(run, size=9.5, color=DARK_GRAY)
                    # Alternate row shading
                    if ri % 2 == 1:
                        set_cell_shading(cell, TABLE_ALT_ROW)

            # Add some space after table
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    # ── Signature block detection and formatting ──
    # The signature lines (Name:, Signature:, Date:) are already handled as paragraphs
    # with bold markers. This looks fine as-is.

    # ── Fix settings.xml zoom element (python-docx bug: missing required percent attr) ──
    settings = doc.settings.element
    zoom = settings.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')

    # ── Save ──
    doc.save(output_path)
    return output_path


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate-policy-docx.py <policy.md> [policy2.md ...]")
        print("Generates a .docx file for each markdown policy file.")
        sys.exit(1)

    for md_path in sys.argv[1:]:
        if not os.path.exists(md_path):
            print(f"File not found: {md_path}")
            continue

        # Output path: same name, .docx extension
        base = os.path.splitext(md_path)[0]
        docx_path = base + '.docx'

        print(f"Parsing: {md_path}")
        blocks = parse_markdown(md_path)

        print(f"Generating: {docx_path}")
        build_docx(blocks, docx_path)

        print(f"Done: {docx_path}")
        print()


if __name__ == '__main__':
    main()
